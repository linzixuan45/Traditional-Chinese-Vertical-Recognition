import math
import os
import re
import sys

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Cm


def get_sort_dir(imgsdir, sort_way=2):
    """

    :param imgsdir:
    :param sort_way: 两种方式，第一种按照整数类型比较大小，第二种按照字符串方式比较大小
    :return:
    """
    img_sort_path = []
    for img in [os.path.join(imgsdir, img) for img in os.listdir(imgsdir)]:

        basename = os.path.basename(img).split('.')[0]
        try:
            basename_int = int(basename)
        except:
            basename_int = 0
        img_sort_path.append([img, basename, basename_int])
    try:
        img_sort_path.sort(key=lambda x: x[sort_way])
    except:
        try:
            if sort_way == 1:
                img_sort_path.sort(key=lambda x: x[2])
            else:
                img_sort_path.sort(key=lambda x: x[1])
        except:
            sys.exit()

    result = np.array(img_sort_path)[:, 0]
    return result


def findnum(string):
    comp=re.compile('-?[1-9]\d*')
    list_str=comp.findall(string)
    list_num=[]
    string = ""
    for item in list_str:
        string = string + item
        item=int(item)
        list_num.append(item)

    return list_num, int(string)


def get_sort_path(imgs_path, sort_way=2):
    img_sort_path = []
    for img in imgs_path:
        basename = os.path.basename(img).split('.')[0]
        try:
            _, basename_int = findnum(basename)
        except:
            basename_int = 0
        img_sort_path.append([img, basename, basename_int])
    img_sort_path.sort(key=lambda x: x[sort_way])

    if img_sort_path:
        result = np.array(img_sort_path)[:, 0]

    else:
        result = img_sort_path

    #  IndexError: too many indices for array: array is 1-dimensional, but 2 were indexed
    # 没有读到文件夹
    return list(result)


def inference_dir(project, imgs_dir, save_dir, img_save_path, dpi, acc_loc=True, task_name="out", sort_way=2, qt_print = None):
    """

    """
    task_name = task_name + f"_acc_{acc_loc}"
    if not os.path.exists(imgs_dir):
        os.makedirs(imgs_dir)
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    if not os.path.exists(img_save_path):
        os.makedirs(img_save_path)

    if not os.path.exists(f"{save_dir}/{task_name}_down_task.csv"):
        with open(f"{save_dir}/{task_name}_down_task.csv",'w') as f:
            f.close()

    imgs_path = get_sort_dir(imgs_dir, sort_way)
    with open(f"{save_dir}/{task_name}_down_task.csv", 'r', encoding='utf-8') as file:
        down_task = file.read().split('\n')
        file.close()

    down_task = list(set(down_task))
    diff_inter_ls = set(imgs_path).difference(set(down_task).intersection(imgs_path))  # 两个列表的并集部分
    diff_inter_ls = get_sort_path(list(diff_inter_ls))
    if qt_print:
        qt_print(diff_inter_ls)
    else:
        print(diff_inter_ls)

    if not diff_inter_ls:
        print(f"img dir ： {imgs_dir}    had already down")
    # 以append的方式不断写入到csv文件中
    """word 文件存在就继续保存，word文件不存在就创建新的对象"""
    if os.path.exists(f"{save_dir}/{task_name}_word.docx"):
        doc1 = Document(f"{save_dir}/{task_name}_word.docx")
        doc2 = Document(f"{save_dir}/{task_name}_img.docx")
    else:
        doc1 = Document()
        doc2 = Document()

    if len(img_save_path) > 100:
        del_file(img_save_path)

    infe_task = []
    with open(f"{save_dir}/{task_name}_json_dict.csv", 'a', encoding='utf8') as name:
        for i, img_path in enumerate(diff_inter_ls):
            res, savepath, json_dict = project.inference(img_path, img_save_path, acc_loc, dpi=dpi)
            doc2.add_picture(savepath, width=Cm(15))


            name.write(str(json_dict) + '\n')

            for word in res:
                doc1.add_paragraph(word)
            doc1.add_paragraph("\n")
            doc1.add_paragraph("\n")
            infe_task.append(img_path)
            if i % 5 == 0:  # 每5首诗保存一次
                doc1.save(f"{save_dir}/{task_name}_word.docx")
                doc2.save(f"{save_dir}/{task_name}_img.docx")
                with open(f"{save_dir}/{task_name}_down_task.csv", 'a+', encoding='utf-8') as file:
                    for path in infe_task:
                        file.write(path)
                        file.write('\n')
                    file.close()
                print_info = "\n".join(infe_task)
                if qt_print:

                    qt_print(f"those imgs have down : \n {print_info}")
                else:
                    print(f"those imgs have down : \n {print_info}")
                infe_task = []
        name.close()  # csv文件关闭


def two_sentences_similarity(sents_1, sents_2):
    counter = 0
    for sent in sents_1:
        if sent in sents_2:
            counter += 1
    try:
        sents_similarity = counter / (math.log(len(sents_1)) + math.log(len(sents_2)))
    except:
        sents_similarity = 0
    return sents_similarity


def str_insert(str_origin, pos, str_add):
    """
    字符串是不可变，只能变为list插入
    """
    str_list = list(str_origin)
    str_list.insert(pos, str_add)
    str_out = "".join(str_list)
    return str_out


def find_insert_sign(origin_sentence, sentence2):
    """

    :param origin_sentence: 原始的准确的句子
    :param sentence2: 相同长度的带有符号的句子
    :return:
    """
    my_re3 = re.compile(r"[-\，\。\？\！\| \… \. \: \（ \） \、 \~ \【 \】 \· \「 \」 \— \/ \{ \} ]", re.S)
    res3 = re.finditer(my_re3, sentence2)
    out_str = origin_sentence
    sign_length = 0
    for i in res3:
        out_str = str_insert(out_str, i.span()[0], i.group())
        sign_length += 1
    if len(sentence2) - sign_length == len(origin_sentence):
        return out_str
    else:
        return origin_sentence


def plt_img_info(img_path, info, save_img_path, dpi=300):
    img = plt.imread(img_path)
    img_name = os.path.basename(img_path)
    plt.subplot(121)
    plt.imshow(img)
    plt.axis('off')

    plt.subplot(122)
    plt.imshow(img)
    for i, loc in enumerate(info['words_result']):
        location = loc['location']
        xy = (location['left'], location['top'])
        h, w = location['height'], location['width']
        plt.gca().add_patch(plt.Rectangle(xy, w, h))
        # plt.Rectangle
        # 第一个参数是坐标(x,y),
        # 第二个参数是矩形宽度 第三个坐标是矩形高度
    plt.axis('off')
    save_path = save_img_path + '/' + img_name
    plt.savefig(save_path, dpi=dpi)
    plt.show()
    plt.close()
    return save_path


def acc_loc_poem(info):
    """
    关于古代繁体的识别xy规则：
    1. 优先比较x大小, x最大为首位
    2. 同x(宽度附近,需要设定阈值), 次选y大小,y小的靠上
    依据获取的文本和位置信息对获取的文本进行排序
    :param info:
    :return: poem_sort, loc_combine  返回排序的位置和需要合并的位置信息
    """
    ocr_loc = []
    for i, loc in enumerate(info['words_result']):
        location = loc['location']
        x, y = location['left'], location['top']
        h, w = location['height'], location['width']
        ocr_loc.append([int(x), int(y), int(w), int(h), int(i)])

    ocr_loc.sort(key=lambda x: (x[0]), reverse=True)
    orc_loc_np = np.array(ocr_loc)

    """ 平均的x阈值  """
    mean_threshold_ls = []
    for i in range(len(orc_loc_np) - 1):
        temp_threshold = orc_loc_np[:, 0][i] - orc_loc_np[:, 0][i + 1]
        mean_threshold_ls.append([temp_threshold, i])
    """去掉最高值和最低值，求取threshold"""



    mean_threshold_ls.sort()
    if len(mean_threshold_ls) >= 3:
        mean_threshold = np.mean(mean_threshold_ls[1:-1])
    else:
        mean_threshold = np.mean(mean_threshold_ls)

    try:  # 杜绝只有一行文字的情况
        mean_w = np.mean(orc_loc_np[:, 2])
    except:
        mean_w = mean_threshold


    threshold = mean_threshold if mean_threshold < mean_w else mean_w
    """如果相减少于阈值，那就是属于同一列，那就对比y值大小,然后合并"""
    try:
        result_poem_sort = list(orc_loc_np[:, 4])  # 最终的诗文排序
    except:
        print(orc_loc_np)
        result_poem_sort = []
    loc_combine = []

    # print(orc_loc_np)
    # print(mean_threshold_ls)
    # print(threshold)
    for i in range(len(orc_loc_np) - 1):
        if orc_loc_np[:, 0][i] - orc_loc_np[:, 0][i + 1] < threshold:  # 交换位置
            loc_combine.append(i)
            if orc_loc_np[:, 1][i] > orc_loc_np[:, 1][i + 1]:  # y的位置,小的在上
                temp = result_poem_sort[i + 1]
                result_poem_sort[i + 1] = result_poem_sort[i]
                result_poem_sort[i] = temp

    """显示排序后的文本"""
    poem_sort = [info['words_result'][i]['words'] for i in result_poem_sort]
    return poem_sort, loc_combine


def get_acc_poem(poem_sort, loc_combine, poem_sample):
    new_poem = []
    for sentence in poem_sort:
        sim_list = []
        for word in poem_sample:
            sim = two_sentences_similarity(word, sentence)
            sim_list.append(sim)
        max_index = 0
        for i, value in enumerate(sim_list):
            if value >= sim_list[max_index]:
                max_index = i
        new_poem.append(find_insert_sign(sentence, poem_sample[max_index]))
    result_poem = []
    jump_flag = False
    for i in range(len(new_poem)):
        if jump_flag:
            jump_flag = False
            continue
        elif i in loc_combine:
            result_poem.append(new_poem[i] + " " + new_poem[i + 1])
            jump_flag = True
        else:
            result_poem.append(new_poem[i])

    return result_poem


def get_poem(poem_sort, loc_combine):
    """
    按照给定的位置合并信息合并
    :param poem_sort:
    :param loc_combine:
    :return:
    """
    result_poem = []
    jump_flag = False
    for i in range(len(poem_sort)):
        if jump_flag:
            jump_flag = False
            continue
        elif i in loc_combine:
            result_poem.append(poem_sort[i] + " " + poem_sort[i + 1])
            jump_flag = True
        else:
            result_poem.append(poem_sort[i])

    return result_poem


# python删除文件的方法 os.remove(path)path指的是文件的绝对路径,如：
def del_file(path_data):
    for i in os.listdir(path_data):  # os.listdir(path_data)#返回一个列表，里面是当前目录下面的所有东西的相对路径
        file_data = path_data + "\\" + i  # 当前文件夹的下面的所有东西的绝对路径
        if os.path.isfile(file_data):  # os.path.isfile判断是否为文件,如果是文件,就删除.如果是文件夹.递归给del_file.
            os.remove(file_data)
        else:
            del_file(file_data)
